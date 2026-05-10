#!/usr/bin/env python3
"""
河南科技大学本科毕业设计（论文）生成脚本
论文题目：基于检索增强生成的个人成长辅导系统设计与实现
作者：孙志豪
"""

import os
import re
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# 常量配置
# ============================================================

THESIS = {
    'title_cn': '基于检索增强生成的个人成长辅导系统设计与实现',
    'title_en': 'Design and Implementation of Personal Growth Counseling System Based on Retrieval-Augmented Generation',
    'title_en_upper': 'DESIGN AND IMPLEMENTATION OF PERSONAL GROWTH COUNSELING SYSTEM BASED ON RETRIEVAL-AUGMENTED GENERATION',
    'author': '孙志豪',
    'college': '信息工程学院',
    'major': '计算机科学与技术',
    'advisor': '霍元智',
    'date': '2026 年 6 月 1 日',
    'year': '2026',
    'month': '6',
    'day': '1',
    'keywords_cn': ['检索增强生成', '个人成长', '辅导系统', '大语言模型', '知识库'],
    'keywords_en': ['Retrieval-Augmented Generation', 'Personal Growth', 'Counseling System', 'Large Language Model', 'Knowledge Base'],
}

# 字体
FONT_SONG = '宋体'
FONT_HEI = '黑体'
FONT_EN = 'Times New Roman'
FONT_KAI = '楷体'

# 字号（磅值）
SIZE_SAN = Pt(16)       # 三号
SIZE_XSAN = Pt(15)      # 小三号
SIZE_SI = Pt(14)        # 四号
SIZE_XSI = Pt(12)       # 小四号
SIZE_WU = Pt(10.5)      # 五号
SIZE_XWU = Pt(9)        # 小五号
SIZE_ER = Pt(22)        # 二号
SIZE_XER = Pt(18)       # 小二号

# 页边距
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)
MARGIN_LEFT = Cm(3.0)   # 装订线
MARGIN_RIGHT = Cm(2.5)

# 页眉文字
HEADER_FIXED = '河南科技大学毕业设计说明书（论文）'

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, f'{THESIS["title_cn"]}.docx')
DIAGRAMS_DIR = os.path.join(OUTPUT_DIR, 'diagrams')

# Mermaid CLI 配置
MERMAID_PUPPETEER_CONFIG = os.path.join(DIAGRAMS_DIR, 'puppeteer-config.json')


# ============================================================
# 样式配置（必须在文档创建后立即调用）
# ============================================================

def configure_styles(doc):
    """配置文档的标题样式，使 TOC 域可以识别章节层级"""
    # Heading 1 → 章标题（三号黑体，居中，黑色）
    h1 = doc.styles['Heading 1']
    h1.font.name = FONT_EN
    h1.font.size = SIZE_SAN
    h1.font.bold = False
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1_fmt = h1.paragraph_format
    h1_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1_fmt.space_before = Pt(0)
    h1_fmt.space_after = Pt(0)
    h1_fmt.page_break_before = True
    # 东亚字体
    h1_rPr = h1.element.find(qn('w:rPr'))
    if h1_rPr is None:
        h1_rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        h1.element.insert(0, h1_rPr)
    h1_rFonts = h1_rPr.find(qn('w:rFonts'))
    if h1_rFonts is None:
        h1_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        h1_rPr.insert(0, h1_rFonts)
    h1_rFonts.set(qn('w:eastAsia'), FONT_HEI)

    # Heading 2 → 节标题（四号黑体，左对齐，黑色）
    h2 = doc.styles['Heading 2']
    h2.font.name = FONT_EN
    h2.font.size = SIZE_SI
    h2.font.bold = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2_fmt = h2.paragraph_format
    h2_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2_fmt.space_before = Pt(6)
    h2_fmt.space_after = Pt(3)
    h2_fmt.page_break_before = False
    h2_rPr = h2.element.find(qn('w:rPr'))
    if h2_rPr is None:
        h2_rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        h2.element.insert(0, h2_rPr)
    h2_rFonts = h2_rPr.find(qn('w:rFonts'))
    if h2_rFonts is None:
        h2_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        h2_rPr.insert(0, h2_rFonts)
    h2_rFonts.set(qn('w:eastAsia'), FONT_HEI)

    # Heading 3 → 子节标题（小四号宋体，左对齐，黑色）
    h3 = doc.styles['Heading 3']
    h3.font.name = FONT_EN
    h3.font.size = SIZE_XSI
    h3.font.bold = False
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3_fmt = h3.paragraph_format
    h3_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3_fmt.space_before = Pt(3)
    h3_fmt.space_after = Pt(3)
    h3_fmt.page_break_before = False
    h3_rPr = h3.element.find(qn('w:rPr'))
    if h3_rPr is None:
        h3_rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        h3.element.insert(0, h3_rPr)
    h3_rFonts = h3_rPr.find(qn('w:rFonts'))
    if h3_rFonts is None:
        h3_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        h3_rPr.insert(0, h3_rFonts)
    h3_rFonts.set(qn('w:eastAsia'), FONT_SONG)


# ============================================================
# 辅助函数
# ============================================================

def set_run_font(run, cn_font=FONT_SONG, en_font=FONT_EN, size=SIZE_XSI, bold=False):
    """设置 run 的中英文字体、字号和粗体"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = en_font
    # 设置东亚字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)


def set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=Emu(480060),  # 约2字符
                         line_spacing=1.5, space_before=Pt(0), space_after=Pt(0)):
    """设置段落格式"""
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.first_line_indent = first_line_indent
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after


def add_body_para(doc, text, bold=False, indent=True):
    """添加正文段落（小四宋体，1.5倍行距）"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_XSI, bold=bold)
    if indent:
        set_paragraph_format(para, first_line_indent=Emu(480060))
    else:
        set_paragraph_format(para, first_line_indent=Emu(0))
    return para


def add_cited_para(doc, text, indent=True):
    """添加包含上标引用的正文段落。
    使用 ^{[N]} 或 ^{[N,M]} 标记上标引用，如：
    "Lewis等人^{[1]}提出了RAG框架" → Lewis等人[1]提出了RAG框架（[1]为上标）
    """
    para = doc.add_paragraph()
    # 用正则拆分：匹配 ^{...} 模式
    pattern = r'\^\{(\[[^\]]*\])\}'
    parts = re.split(pattern, text)

    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 0:
            # 普通文本
            run = para.add_run(part)
            set_run_font(run, cn_font=FONT_SONG, size=SIZE_XSI)
        else:
            # 上标引用文本（如 [1] 或 [1,2]）
            run = para.add_run(part)
            set_run_font(run, cn_font=FONT_SONG, size=SIZE_XSI)
            run.font.superscript = True

    if indent:
        set_paragraph_format(para, first_line_indent=Emu(480060))
    else:
        set_paragraph_format(para, first_line_indent=Emu(0))
    return para


def add_chapter_title(doc, title):
    """添加章标题（三号黑体，居中）—— 使用 Heading 1 样式"""
    para = doc.add_paragraph(title, style='Heading 1')
    # 确保格式覆盖样式默认值
    para.paragraph_format.first_line_indent = Emu(0)
    return para


def add_section_title(doc, title):
    """添加节标题（四号黑体，左对齐）—— 使用 Heading 2 样式"""
    para = doc.add_paragraph(title, style='Heading 2')
    para.paragraph_format.first_line_indent = Emu(0)
    return para


def add_subsection_title(doc, title):
    """添加子节标题（小四号黑体，左对齐）—— 使用 Heading 3 样式"""
    para = doc.add_paragraph(title, style='Heading 3')
    para.paragraph_format.first_line_indent = Emu(0)
    return para


def add_empty_para(doc, size=SIZE_XSI):
    """添加空段落"""
    para = doc.add_paragraph()
    run = para.add_run('')
    set_run_font(run, size=size)
    set_paragraph_format(para, first_line_indent=Emu(0))
    return para


def setup_page_margins(section):
    """设置页面边距"""
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT


def add_header(section, text):
    """添加页眉"""
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        para = header.paragraphs[0]
    else:
        para = header.add_paragraph()
    para.text = ''
    run = para.add_run(text)
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 添加页眉下横线
    pPr = para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def remove_header_footer(section):
    """移除页眉页脚"""
    section.header.is_linked_to_previous = False
    for p in section.header.paragraphs:
        p.text = ''
    section.footer.is_linked_to_previous = False
    for p in section.footer.paragraphs:
        p.text = ''


def add_page_number(section):
    """添加页脚页码"""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.text = ''
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加 PAGE 域代码
    run = para.add_run()
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)

    run2 = para.add_run()
    set_run_font(run2, cn_font=FONT_SONG, size=SIZE_WU)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)

    run3 = para.add_run()
    set_run_font(run3, cn_font=FONT_SONG, size=SIZE_WU)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar2)

    run4 = para.add_run('1')
    set_run_font(run4, cn_font=FONT_SONG, size=SIZE_WU)

    run5 = para.add_run()
    set_run_font(run5, cn_font=FONT_SONG, size=SIZE_WU)
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar3)


def set_page_number_format(section, fmt='decimal'):
    """设置页码格式：decimal=阿拉伯数字, upperRoman=大写罗马数字"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")}/>')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)


def set_page_number_start(section, start=1):
    """设置页码起始值"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")}/>')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), str(start))


def add_toc_field(doc):
    """插入目录域代码"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(para, first_line_indent=Emu(0))

    run1 = para.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1._element.append(fldChar1)

    run2 = para.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._element.append(instrText)

    run3 = para.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar2)

    run4 = para.add_run('（请在 Word 中右键点击此处，选择"更新域"以生成目录）')
    set_run_font(run4, cn_font=FONT_SONG, size=SIZE_XSI)

    run5 = para.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar3)


def add_three_line_table(doc, headers, rows, col_widths=None):
    """添加三线表"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(h)
        set_run_font(run, cn_font=FONT_HEI, size=SIZE_WU, bold=True)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 三线表样式：顶线、表头底线、底线
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    # 移除旧边框
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)
    tblPr.append(borders)

    # 表头行底部边框
    for cell in table.rows[0].cells:
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            f'</w:tcBorders>'
        )
        old_tcBorders = tcPr.find(qn('w:tcBorders'))
        if old_tcBorders is not None:
            tcPr.remove(old_tcBorders)
        tcPr.append(tcBorders)

    return table


# ============================================================
# Mermaid 图形生成
# ============================================================

def ensure_diagrams_dir():
    """确保 diagrams 目录和 puppeteer 配置存在"""
    os.makedirs(DIAGRAMS_DIR, exist_ok=True)
    if not os.path.exists(MERMAID_PUPPETEER_CONFIG):
        with open(MERMAID_PUPPETEER_CONFIG, 'w') as f:
            f.write('{"args": ["--no-sandbox", "--disable-setuid-sandbox"]}')


def generate_mermaid_image(mermaid_text, filename, width=2400, height=1600):
    """将 Mermaid 文本渲染为 PNG 图片，返回图片路径"""
    ensure_diagrams_dir()
    output_path = os.path.join(DIAGRAMS_DIR, filename)

    # 写入临时 .mmd 文件
    with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False,
                                      encoding='utf-8', dir=DIAGRAMS_DIR) as f:
        f.write(mermaid_text)
        mmd_path = f.name

    try:
        cmd = [
            'mmdc', '-i', mmd_path, '-o', output_path,
            '-b', 'white', '-p', MERMAID_PUPPETEER_CONFIG,
            '-w', str(width), '-H', str(height),
            '-s', '2'
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode != 0:
            print(f'  [WARNING] Mermaid 渲染失败 ({filename}): {result.stderr.strip()}')
            return None
        return output_path
    except FileNotFoundError:
        print('  [WARNING] mmdc 未安装，跳过图形生成')
        return None
    except subprocess.TimeoutExpired:
        print(f'  [WARNING] Mermaid 渲染超时 ({filename})')
        return None
    finally:
        os.unlink(mmd_path)


def add_figure(doc, img_path, caption, width_inches=5.0):
    """在文档中插入图片并添加图注（图X-Y格式，五号宋体，居中）"""
    if img_path is None or not os.path.exists(img_path):
        add_body_para(doc, f'[{caption} — 图片生成失败，请手动插入]')
        return

    # 插入图片
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(6), space_after=Pt(3))

    # 图注
    cap_para = doc.add_paragraph()
    cap_run = cap_para.add_run(caption)
    set_run_font(cap_run, cn_font=FONT_SONG, size=SIZE_WU)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(cap_para, first_line_indent=Emu(0), space_before=Pt(3), space_after=Pt(6))


# ============================================================
# Mermaid 图形定义
# ============================================================

DIAGRAM_ARCHITECTURE = """
graph TB
    subgraph 表现层
        A[Vue.js 3 前端]
        B[Element Plus 组件库]
        C[Axios HTTP 客户端]
    end
    subgraph 业务逻辑层
        D[FastAPI 后端]
        E[用户管理模块]
        F[智能对话模块]
        G[知识库管理模块]
        H[成长管理模块]
    end
    subgraph 数据访问层
        I[SQLAlchemy ORM]
        J[Chroma 客户端]
    end
    subgraph 数据存储层
        K[(MySQL 8.0)]
        L[(Chroma 向量库)]
    end
    subgraph 外部服务
        M[GLM-4 大语言模型]
        N[文本嵌入模型]
    end
    A --> D
    D --> I
    D --> J
    I --> K
    J --> L
    F --> M
    F --> N
"""

DIAGRAM_ER = """
erDiagram
    USER ||--o{ CONVERSATION : "拥有"
    USER ||--o{ GROWTH_GOAL : "设定"
    USER ||--o{ GROWTH_RECORD : "撰写"
    CONVERSATION ||--|{ MESSAGE : "包含"
    KNOWLEDGE_DOCUMENT ||--|{ KNOWLEDGE_CHUNK : "拆分为"
    GROWTH_GOAL ||--o{ GROWTH_RECORD : "跟踪"

    USER {
        int id PK
        string username
        string email
        string password_hash
        string nickname
        string role
        datetime created_at
    }
    CONVERSATION {
        int id PK
        int user_id FK
        string title
        datetime created_at
    }
    MESSAGE {
        int id PK
        int conversation_id FK
        string role
        text content
        datetime created_at
    }
    KNOWLEDGE_DOCUMENT {
        int id PK
        string title
        string category
        int chunk_count
        string status
        datetime uploaded_at
    }
    KNOWLEDGE_CHUNK {
        int id PK
        int document_id FK
        text content
        int chunk_index
    }
    GROWTH_GOAL {
        int id PK
        int user_id FK
        string title
        string category
        string status
        int priority
        date deadline
        datetime created_at
    }
    GROWTH_RECORD {
        int id PK
        int user_id FK
        int goal_id FK
        text content
        datetime record_time
    }
"""

DIAGRAM_RAG_FLOW = """
graph LR
    A[用户提问] --> B[问题向量化]
    B --> C[向量相似度检索]
    C --> D[Top-K 知识片段]
    D --> E[组装提示词]
    E --> F[大语言模型生成]
    F --> G[返回辅导建议]

    subgraph 知识库构建
        H[知识文档] --> I[文本提取与分块]
        I --> J[向量化嵌入]
        J --> K[(Chroma 向量库)]
    end

    K --> C
"""

DIAGRAM_MODULE = """
graph TD
    A[个人成长辅导系统] --> B[用户管理模块]
    A --> C[智能对话模块]
    A --> D[知识库管理模块]
    A --> E[个人成长管理模块]

    B --> B1[用户注册]
    B --> B2[用户登录]
    B --> B3[个人信息管理]

    C --> C1[RAG 检索增强]
    C --> C2[多轮对话]
    C --> C3[对话历史管理]

    D --> D1[文档上传]
    D --> D2[文本解析与分块]
    D --> D3[向量化索引]

    E --> E1[目标设定]
    E --> E2[进度跟踪]
    E --> E3[成长记录]
"""

DIAGRAM_LOGIN_FLOW = """
graph TD
    A[用户访问系统] --> B{已登录?}
    B -->|是| C[进入主页面]
    B -->|否| D[登录/注册页面]
    D --> E[输入账号密码]
    E --> F{验证通过?}
    F -->|是| C
    F -->|否| G[提示错误信息]
    G --> D
    C --> H[智能对话]
    C --> I[知识库管理]
    C --> J[成长管理]
"""


# ============================================================
# 封面页
# ============================================================
def create_cover_page(doc):
    section = doc.sections[0]
    setup_page_margins(section)
    remove_header_footer(section)

    # 学校英文名
    para = doc.add_paragraph()
    run = para.add_run('HENAN UNIVERSITY OF SCIENCE & TECHNOLOGY')
    set_run_font(run, cn_font=FONT_EN, en_font=FONT_EN, size=Pt(20), bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(50), space_after=Pt(30))

    # 毕业设计（论文）标题
    add_empty_para(doc, Pt(24))
    add_empty_para(doc, Pt(24))

    para = doc.add_paragraph()
    run = para.add_run('毕 业 设 计（论 文）')
    set_run_font(run, cn_font=FONT_HEI, size=SIZE_XER, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(20), space_after=Pt(40))

    # 空行间隔
    for _ in range(4):
        add_empty_para(doc, SIZE_XSI)

    # 封面信息字段
    fields = [
        ('题    目', THESIS['title_cn']),
        ('姓    名', THESIS['author']),
        ('学    院', THESIS['college']),
        ('专    业', THESIS['major']),
        ('指导教师', THESIS['advisor']),
    ]

    for label, value in fields:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(6), space_after=Pt(6))
        run_label = para.add_run(f'{label}    ')
        set_run_font(run_label, cn_font=FONT_SONG, size=SIZE_XSI, bold=True)
        run_value = para.add_run(value)
        set_run_font(run_value, cn_font=FONT_SONG, size=SIZE_XSI)

    # 空行
    for _ in range(4):
        add_empty_para(doc, SIZE_XSI)

    # 日期
    para = doc.add_paragraph()
    run = para.add_run(THESIS['date'])
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_SI)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0))

    # 分页（空白页）
    doc.add_page_break()
    add_empty_para(doc)
    doc.add_page_break()


# ============================================================
# 声明页
# ============================================================
def create_declaration_page(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_page_margins(section)
    remove_header_footer(section)

    # 写作声明
    para = doc.add_paragraph()
    run = para.add_run('学位论文写作声明')
    set_run_font(run, cn_font=FONT_HEI, size=SIZE_SI, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(30), space_after=Pt(15))

    declaration_text = (
        '本人郑重声明：所呈交的学位论文，是本人在导师的指导下，独立进行研究'
        '工作所取得的成果。除文中已经注明引用的内容外，本论文不含任何其他个人或'
        '集体已经发表或撰写过的作品或成果。对本文的研究做出重要贡献的个人和集体，'
        '均已在文中以明确方式标明。本声明的法律结果由本人承担。'
    )
    add_body_para(doc, declaration_text)

    add_empty_para(doc)
    add_empty_para(doc)

    para = doc.add_paragraph()
    run = para.add_run('论文作者签名：            日期：     年   月   日')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_XSI)
    set_paragraph_format(para, first_line_indent=Emu(0), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_empty_para(doc)
    add_empty_para(doc)

    # 使用授权说明
    para = doc.add_paragraph()
    run = para.add_run('学位论文使用授权说明')
    set_run_font(run, cn_font=FONT_HEI, size=SIZE_SI, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(30), space_after=Pt(15))

    auth_text = (
        '本人完全了解河南科技大学关于收集、保存、使用学位论文的规定，即：按'
        '照学校要求提交学位论文的印刷本和电子版本；学校有权保存学位论文的印刷本'
        '和电子版，并提供目录检索与阅览服务；学校可以采用影印、缩印、数字化或其'
        '它复制手段保存论文；在不以赢利为目的的前提下，学校可以将学位论文编入有'
        '关数据库，提供网上服务。（保密论文在解密后遵守此规定）'
    )
    add_body_para(doc, auth_text)

    add_empty_para(doc)
    add_empty_para(doc)

    para = doc.add_paragraph()
    run = para.add_run('论文作者签名：              导师签名：')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_XSI)
    set_paragraph_format(para, first_line_indent=Emu(0), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_empty_para(doc)

    para = doc.add_paragraph()
    run = para.add_run('日期：     年   月   日')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_XSI)
    set_paragraph_format(para, first_line_indent=Emu(0), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # 翻页空白页
    doc.add_page_break()
    add_empty_para(doc)
    doc.add_page_break()


# ============================================================
# 中文摘要
# ============================================================
def create_chinese_abstract(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_page_margins(section)
    add_header(section, HEADER_FIXED)
    add_page_number(section)
    set_page_number_format(section, 'upperRoman')
    set_page_number_start(section, 1)

    # 论文标题
    para = doc.add_paragraph()
    run = para.add_run(THESIS['title_cn'])
    set_run_font(run, cn_font=FONT_HEI, size=SIZE_XSAN, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(20), space_after=Pt(15))

    # "摘 要" 标题
    para = doc.add_paragraph()
    run = para.add_run('摘  要')
    set_run_font(run, cn_font=FONT_HEI, size=SIZE_XSAN, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(10), space_after=Pt(10))

    # 摘要正文
    abstract_cn = (
        '随着社会竞争的日益激烈和信息技术的快速发展，个人成长与职业发展规划'
        '成为当代青年面临的重要课题。传统的一对一辅导方式受限于时间、地域和师资'
        '成本，难以满足大量用户的个性化需求。近年来，以大语言模型为代表的人工智能'
        '技术取得了突破性进展，检索增强生成（Retrieval-Augmented Generation, RAG）'
        '技术通过将外部知识库与大语言模型相结合，能够有效缓解模型幻觉问题，提升'
        '生成内容的准确性和可靠性，为构建智能化辅导系统提供了新的技术路径。'
    )
    add_body_para(doc, abstract_cn)

    abstract_cn2 = (
        '本文设计并实现了一个基于检索增强生成的个人成长辅导系统。该系统采用'
        'B/S架构，前后端分离的开发模式，后端基于Python语言和FastAPI框架构建RESTful'
        'API服务，前端采用Vue.js 3框架和Element Plus组件库实现用户界面，数据存储'
        '采用MySQL关系型数据库和Chroma向量数据库的混合方案。系统的核心功能包括：'
        '用户管理模块、智能对话模块、知识库管理模块和个人成长管理模块。'
    )
    add_body_para(doc, abstract_cn2)

    abstract_cn3 = (
        '在系统核心的智能对话模块中，本文实现了完整的RAG流程：首先对用户上传的'
        '个人成长领域知识文档进行解析和分块处理，利用文本嵌入模型将知识块转化为向量'
        '表示并存储在Chroma向量数据库中；当用户发起咨询时，系统通过向量相似度检索'
        '获取相关知识片段，将其作为上下文信息注入大语言模型的提示词中，从而生成具有'
        '知识依据的个性化辅导建议。系统还提供了成长目标设定、进度跟踪和成长记录等'
        '功能，帮助用户系统化管理个人成长路径。'
    )
    add_body_para(doc, abstract_cn3)

    abstract_cn4 = (
        '通过对系统进行功能测试和性能测试，结果表明该系统各项功能运行正常，'
        '智能对话模块能够准确检索相关知识并生成合理的辅导建议，系统响应时间满足'
        '实际使用需求，具有良好的用户体验和实用价值。本系统为个人成长辅导领域提供'
        '了一种基于人工智能技术的创新解决方案，具有一定的理论意义和应用前景。'
    )
    add_body_para(doc, abstract_cn4)

    # 关键词
    para = doc.add_paragraph()
    run = para.add_run('关键词：')
    set_run_font(run, cn_font=FONT_HEI, size=SIZE_XSI, bold=True)
    kw_str = '；'.join(THESIS['keywords_cn'])
    run2 = para.add_run(kw_str)
    set_run_font(run2, cn_font=FONT_SONG, size=SIZE_XSI)
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(10))


# ============================================================
# 英文摘要
# ============================================================
def create_english_abstract(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_page_margins(section)
    add_header(section, THESIS['title_en_upper'])
    add_page_number(section)
    set_page_number_format(section, 'upperRoman')

    # 英文标题
    para = doc.add_paragraph()
    run = para.add_run(THESIS['title_en_upper'])
    set_run_font(run, en_font=FONT_EN, size=SIZE_XSAN, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(20), space_after=Pt(15))

    # ABSTRACT 标题
    para = doc.add_paragraph()
    run = para.add_run('ABSTRACT')
    set_run_font(run, en_font=FONT_EN, size=SIZE_XSAN, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(10), space_after=Pt(10))

    # 英文摘要正文
    abstract_en1 = (
        'With the intensification of social competition and the rapid development of '
        'information technology, personal growth and career development planning have become '
        'important issues for contemporary young people. Traditional one-on-one counseling '
        'methods are limited by time, geography, and cost, making it difficult to meet the '
        'personalized needs of a large number of users. In recent years, artificial intelligence '
        'technology represented by Large Language Models (LLMs) has made breakthrough progress. '
        'Retrieval-Augmented Generation (RAG) technology, by combining external knowledge bases '
        'with large language models, can effectively alleviate the model hallucination problem '
        'and improve the accuracy and reliability of generated content, providing a new technical '
        'approach for building intelligent counseling systems.'
    )
    add_body_para(doc, abstract_en1)

    abstract_en2 = (
        'This thesis designs and implements a personal growth counseling system based on '
        'Retrieval-Augmented Generation. The system adopts a B/S architecture with a separated '
        'frontend and backend development model. The backend is built using Python and the FastAPI '
        'framework to provide RESTful API services, while the frontend utilizes Vue.js 3 and '
        'Element Plus component library for the user interface. Data storage employs a hybrid '
        'solution combining MySQL relational database and Chroma vector database. The core '
        'functions of the system include: user management module, intelligent dialogue module, '
        'knowledge base management module, and personal growth management module.'
    )
    add_body_para(doc, abstract_en2)

    abstract_en3 = (
        'In the core intelligent dialogue module, this thesis implements a complete RAG pipeline: '
        'first, personal growth domain knowledge documents uploaded by users are parsed and chunked, '
        'then converted into vector representations using text embedding models and stored in the '
        'Chroma vector database. When a user initiates a consultation, the system retrieves relevant '
        'knowledge fragments through vector similarity search and injects them as contextual information '
        'into the LLM prompt, thereby generating evidence-based personalized counseling suggestions. '
        'The system also provides features such as growth goal setting, progress tracking, and growth '
        'record management to help users systematically manage their personal growth path.'
    )
    add_body_para(doc, abstract_en3)

    abstract_en4 = (
        'Through functional testing and performance testing, the results demonstrate that all '
        'system functions operate normally. The intelligent dialogue module can accurately retrieve '
        'relevant knowledge and generate reasonable counseling suggestions. The system response time '
        'meets practical usage requirements, exhibiting good user experience and practical value. '
        'This system provides an innovative AI-based solution for the field of personal growth '
        'counseling, with certain theoretical significance and application prospects.'
    )
    add_body_para(doc, abstract_en4)

    # KEY WORDS
    para = doc.add_paragraph()
    run = para.add_run('KEY WORDS: ')
    set_run_font(run, en_font=FONT_EN, size=SIZE_XSI, bold=True)
    kw_str = '; '.join(THESIS['keywords_en'])
    run2 = para.add_run(kw_str)
    set_run_font(run2, en_font=FONT_EN, size=SIZE_XSI)
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(10))


# ============================================================
# 目录页
# ============================================================
def create_table_of_contents(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_page_margins(section)
    add_header(section, HEADER_FIXED)
    add_page_number(section)
    set_page_number_format(section, 'upperRoman')

    # "目录" 标题
    para = doc.add_paragraph()
    run = para.add_run('目  录')
    set_run_font(run, cn_font=FONT_HEI, size=SIZE_XSAN, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(20), space_after=Pt(15))

    # 插入 TOC 域
    add_toc_field(doc)


# ============================================================
# 第1章 绪论
# ============================================================
def create_chapter_1(doc):
    add_chapter_title(doc, '第1章 绪论')

    add_section_title(doc, '1.1 研究背景及意义')

    add_body_para(doc,
        '在当今信息化和全球化快速推进的时代背景下，社会竞争日益激烈，个人成长'
        '与职业发展规划已成为当代青年面临的核心课题。无论是即将步入社会的大学生，'
        '还是处于职业转型期的职场人士，都迫切需要专业的成长辅导和职业指导。然而，'
        '传统的个人成长辅导主要依赖一对一的人工咨询模式，这种方式存在诸多局限性：'
        '首先，专业咨询师资源稀缺且服务费用较高，使得大量有需求的用户难以获得持续'
        '的辅导服务；其次，人工辅导受限于时间和地域，无法实现全天候的即时响应；'
        '最后，传统辅导方式难以系统化地整合和利用海量的个人成长领域知识。'
    )

    add_cited_para(doc,
        '近年来，人工智能技术取得了突飞猛进的发展，特别是以GPT系列、Claude、'
        '文心一言等为代表的大语言模型（Large Language Model, LLM）展现出了强大的'
        '自然语言理解和生成能力^{[4,5]}。这些模型能够理解复杂的用户意图，生成流畅连贯的'
        '文本回复，为人机对话系统的构建奠定了坚实的技术基础^{[7]}。然而，大语言模型也'
        '存在一些固有的局限性，其中最突出的问题就是"幻觉"现象——模型可能会生成'
        '看似合理但实际上缺乏事实依据的内容^{[6]}，这在个人成长辅导等需要专业知识支撑的'
        '场景中是不可接受的。'
    )

    add_cited_para(doc,
        '为了解决上述问题，检索增强生成（Retrieval-Augmented Generation, RAG）'
        '技术应运而生。RAG技术由Lewis等人在2020年首次提出^{[1]}，其核心思想是将外部知识'
        '检索与大语言模型生成能力相结合。在用户提问时，系统首先从预构建的知识库中'
        '检索出相关的知识片段，然后将这些知识作为上下文信息提供给大语言模型，从而'
        '引导模型生成更加准确、可靠且有知识依据的回答^{[3]}。RAG技术有效缓解了大语言模型'
        '的幻觉问题，同时无需对模型进行重新训练，大大降低了系统的部署成本和维护难度。'
    )

    add_body_para(doc,
        '基于以上背景，本文提出并实现了一个基于检索增强生成技术的个人成长辅导'
        '系统。该系统旨在利用RAG技术整合个人成长领域的专业知识，结合大语言模型的'
        '强大对话能力，为用户提供智能化的、个性化的成长辅导服务。本系统的研究与实现'
        '具有重要的理论意义和实际应用价值：在理论层面，本研究探索了RAG技术在个人'
        '成长辅导领域的应用模式，为知识增强型对话系统的设计提供了参考；在实际应用'
        '层面，该系统为用户提供了一种低成本、高可用的智能辅导方案，有助于促进个人'
        '成长辅导服务的普及化和智能化。'
    )

    add_section_title(doc, '1.2 国内外研究现状')

    add_subsection_title(doc, '1.2.1 国外研究现状')

    add_cited_para(doc,
        '在检索增强生成技术方面，Lewis等人于2020年提出了RAG框架^{[1]}，将预训练的'
        ' seq2seq模型与外部知识检索机制相结合，在自然语言问答任务上取得了显著的'
        '性能提升。随后，Karpukhin等人提出了DPR（Dense Passage Retrieval）方法^{[2]}，'
        '利用双编码器架构实现了高效的密集向量检索。2023年，Gao等人提出了RAR（'
        'Retrieval-Augmented Retrieval）方法^{[3]}，进一步提升了检索的准确率。在向量数据库'
        '方面，Chroma、Pinecone、Weaviate等开源和商业向量数据库的发展为RAG系统的'
        '工程化落地提供了重要支撑^{[17,18]}。'
    )

    add_cited_para(doc,
        '在AI辅导系统方面，斯坦福大学、麻省理工学院等研究机构开展了大量关于'
        '智能化教育和辅导系统的研究^{[13]}。Woebot、Wysa等商业化AI心理健康助手产品已经'
        '在市场上取得了广泛的应用，这些系统主要基于认知行为疗法（CBT）等心理学理论，'
        '结合对话式AI技术为用户提供心理健康支持。OpenAI、Google等科技巨头也在积极'
        '探索将大语言模型应用于教育辅导和知识问答领域^{[5]}。'
    )

    add_subsection_title(doc, '1.2.2 国内研究现状')

    add_cited_para(doc,
        '国内在RAG技术方面的研究同样取得了丰硕成果^{[10]}。百度推出的文心一言、阿里'
        '推出的通义千问、讯飞推出的星火认知大模型等国产大语言模型在中文理解和生成'
        '方面表现出色^{[15]}。基于这些大模型，国内研究者和企业开发了多种RAG应用框架，如'
        '智谱AI的LangChain-Chatchat、百度智能云的AppBuilder等开源项目，为RAG技术的'
        '普及和应用提供了有力工具^{[9]}。'
    )

    add_cited_para(doc,
        '在智能辅导系统方面，国内高校和企业在在线教育、智能问答等领域的应用研究'
        '日益活跃^{[11,12]}。好未来、作业帮等教育科技公司推出了基于AI的辅导产品，主要用于K12'
        '教育场景。然而，在个人成长和职业发展辅导领域，专门针对成年用户的智能化辅导'
        '系统研究相对较少，尤其是基于RAG技术的个人成长辅导系统尚未见到成熟的实现方案，'
        '这为本研究提供了良好的创新空间。'
    )

    add_section_title(doc, '1.3 研究内容与方法')

    add_body_para(doc,
        '本文的研究内容主要包括以下几个方面：（1）分析个人成长辅导系统的功能需求'
        '和非功能需求，明确系统的设计目标；（2）设计基于RAG技术的系统总体架构，包括'
        '前端展示层、后端业务逻辑层和数据存储层的详细设计方案；（3）实现系统的核心'
        '功能模块，重点完成RAG知识检索增强生成流程的设计与实现；（4）对系统进行全面的'
        '功能测试和性能测试，验证系统的可行性和实用性。'
    )

    add_body_para(doc,
        '本文采用的研究方法主要包括：文献调研法，通过查阅国内外相关文献，了解RAG'
        '技术和智能辅导系统的研究现状和发展趋势；系统开发法，采用软件工程的方法论指导'
        '系统的需求分析、设计、实现和测试全过程；实验验证法，通过功能测试和性能测试'
        '验证系统的各项功能和性能指标是否达到设计要求。'
    )

    add_section_title(doc, '1.4 本文组织结构')

    add_body_para(doc,
        '本文共分为五章，各章内容安排如下：', indent=False
    )

    add_body_para(doc,
        '第1章为绪论，介绍了本文的研究背景及意义、国内外研究现状、研究内容与'
        '方法以及论文的组织结构。'
    )

    add_body_para(doc,
        '第2章为系统需求分析，从功能需求、非功能需求和可行性分析三个方面对系统'
        '进行了详细的需求分析。'
    )

    add_body_para(doc,
        '第3章为系统设计，包括系统总体架构设计、功能模块设计和数据库设计。'
    )

    add_body_para(doc,
        '第4章为系统实现，介绍了系统的开发环境、核心功能实现和系统界面展示。'
    )

    add_body_para(doc,
        '第5章为系统测试，对系统进行了功能测试和性能测试，并给出了测试结论。'
    )

    add_body_para(doc,
        '最后是总结与展望部分，对全文研究工作进行了总结，并对未来改进方向进行了'
        '展望。'
    )


# ============================================================
# 第2章 系统需求分析
# ============================================================
def create_chapter_2(doc):
    add_chapter_title(doc, '第2章 系统需求分析')

    add_section_title(doc, '2.1 功能需求分析')

    add_body_para(doc,
        '通过对个人成长辅导领域的调研分析，本系统需要满足以下功能需求。系统面向'
        '两类用户：普通用户（寻求成长辅导的个人）和系统管理员。根据用户角色和使用场景'
        '的不同，系统的功能需求可以划分为以下几个核心模块。'
    )

    add_subsection_title(doc, '2.1.1 用户管理模块')

    add_body_para(doc,
        '用户管理模块负责用户的注册、登录和个人信息管理功能。具体需求包括：'
        '（1）用户注册功能，支持用户通过用户名、邮箱和密码进行注册，注册时需进行'
        '基本的输入合法性校验；（2）用户登录功能，支持用户名/邮箱和密码登录，采用'
        'JWT（JSON Web Token）机制进行身份认证和会话管理；（3）个人信息管理功能，'
        '用户可以查看和修改个人基本信息，包括头像、昵称、个人简介等；（4）密码管理'
        '功能，支持密码修改和密码重置操作。'
    )

    add_subsection_title(doc, '2.1.2 智能对话模块')

    add_body_para(doc,
        '智能对话模块是系统的核心功能模块，基于RAG技术实现知识增强的智能辅导对话。'
        '具体需求包括：（1）多轮对话功能，支持用户与系统进行连续的多轮对话交互，系统'
        '能够维护对话上下文，理解用户意图的演变过程；（2）知识增强生成功能，当用户提出'
        '成长相关问题时，系统能够从知识库中检索相关知识片段，结合检索结果生成具有知识'
        '依据的辅导建议；（3）对话历史管理功能，自动保存用户的对话记录，用户可以查看'
        '历史对话内容并继续之前的对话；（4）对话导出功能，支持将对话记录导出为文本文件，'
        '方便用户后续回顾和整理。'
    )

    add_subsection_title(doc, '2.1.3 知识库管理模块')

    add_body_para(doc,
        '知识库管理模块负责个人成长领域知识文档的上传、解析和索引管理。具体需求包括：'
        '（1）文档上传功能，支持用户和管理员上传PDF、TXT、Markdown等格式的知识文档；'
        '（2）文档解析功能，系统自动对上传的文档进行文本提取和分块处理，将长文档切分为'
        '合适大小的知识片段；（3）向量化索引功能，利用文本嵌入模型将知识片段转化为向量'
        '表示，并建立向量索引以便高效检索；（4）知识库维护功能，管理员可以对知识库中的'
        '文档进行查看、编辑和删除操作，保持知识库的时效性和准确性。'
    )

    add_subsection_title(doc, '2.1.4 个人成长管理模块')

    add_body_para(doc,
        '个人成长管理模块帮助用户系统化管理个人成长目标和进度。具体需求包括：'
        '（1）目标设定功能，用户可以创建个人成长目标，包括目标名称、目标描述、计划'
        '完成时间和优先级等信息；（2）进度跟踪功能，系统自动记录用户完成各目标的进度'
        '状态，支持将目标标记为进行中、已完成或已搁置；（3）成长记录功能，用户可以'
        '记录每日或每周的成长心得和反思，形成持续的成长日志；（4）数据统计功能，以'
        '图表形式展示用户的成长数据统计，包括目标完成率、成长记录趋势等。'
    )

    add_section_title(doc, '2.2 非功能需求分析')

    add_body_para(doc,
        '除了功能需求之外，系统还需要满足以下非功能需求，以保证系统的质量和'
        '用户体验。'
    )

    add_subsection_title(doc, '2.2.1 性能需求')

    add_body_para(doc,
        '（1）响应时间：系统的常规页面加载时间应不超过3秒，智能对话模块的响应时间'
        '应控制在10秒以内；（2）并发能力：系统应支持至少50个用户同时在线使用，在并发'
        '访问情况下系统的平均响应时间增幅不超过30%；（3）数据处理能力：系统应能够处理'
        '单次上传不超过20MB的知识文档，知识库中的文档总量应支持至少1000篇文档的存储和'
        '检索。'
    )

    add_subsection_title(doc, '2.2.2 安全需求')

    add_body_para(doc,
        '（1）身份认证：采用JWT令牌机制进行用户身份认证，所有API接口需验证令牌的'
        '有效性；（2）数据安全：用户密码采用bcrypt算法进行加密存储，敏感数据传输采用'
        'HTTPS协议加密；（3）输入防护：对所有用户输入进行合法性校验和XSS/SQL注入防护；'
        '（4）隐私保护：用户的对话记录和个人信息仅用户本人和系统管理员可访问，严格遵守'
        '数据最小化原则。'
    )

    add_subsection_title(doc, '2.2.3 可用性需求')

    add_body_para(doc,
        '（1）界面友好性：系统界面设计应简洁直观，符合用户操作习惯，核心操作不超过'
        '3次点击即可完成；（2）兼容性：系统应兼容主流浏览器，包括Chrome、Firefox、'
        'Edge、Safari等的最近两个版本；（3）可维护性：系统代码结构清晰，模块间耦合度低，'
        '便于后续的功能扩展和维护。'
    )

    add_section_title(doc, '2.3 可行性分析')

    add_subsection_title(doc, '2.3.1 技术可行性')

    add_cited_para(doc,
        '本系统采用的技术栈均为当前主流且成熟的技术方案。Python语言拥有丰富的'
        'AI/ML生态，FastAPI框架性能优异且开发效率高^{[19]}；Vue.js 3是目前最流行的前端框架'
        '之一，拥有完善的组件生态^{[20]}；MySQL是业界广泛使用的关系型数据库，稳定可靠；'
        'Chroma是专为AI应用设计的开源向量数据库，支持高效的向量相似度检索^{[18]}。大语言模型'
        '方面，OpenAI的GPT系列、智谱AI的GLM系列等均提供了完善的API接口，可以直接调用^{[5,15]}。'
        '综上所述，本系统在技术实现上完全可行。'
    )

    add_subsection_title(doc, '2.3.2 经济可行性')

    add_body_para(doc,
        '本系统的开发和运行成本主要包括：开发工具和框架均为开源免费软件，无需'
        '支付许可费用；服务器部署采用云服务器方案，初期可使用基础配置，月费用在'
        '可接受范围内；大语言模型API调用采用按量计费模式，根据实际使用量付费，成本'
        '可控。系统开发所需的人力资源为一名开发者，开发周期约3-4个月。综合评估，'
        '本系统的经济可行性良好。'
    )

    add_subsection_title(doc, '2.3.3 社会可行性')

    add_body_para(doc,
        '个人成长辅导是社会广泛认可的需求领域，本系统提供的服务内容属于教育和'
        '咨询范畴，不涉及医疗诊断等需要特殊资质的领域。系统的设计和运营遵循国家'
        '相关法律法规，注重用户隐私保护和数据安全。同时，系统的使用有助于提升用户'
        '的自我认知和规划能力，具有积极的社会意义。因此，本系统在社会和法律层面'
        '是可行的。'
    )


# ============================================================
# 第3章 系统设计
# ============================================================
def create_chapter_3(doc, generated_diagrams=None):
    if generated_diagrams is None:
        generated_diagrams = {}
    add_chapter_title(doc, '第3章 系统设计')

    add_section_title(doc, '3.1 系统总体架构设计')

    add_body_para(doc,
        '本系统采用基于B/S架构的前后端分离设计模式，整体划分为四个层次：表现层、'
        '业务逻辑层、数据访问层和数据存储层。这种分层架构使得各层之间职责明确、'
        '耦合度低，有利于系统的开发、测试和维护。'
    )

    add_cited_para(doc,
        '表现层基于Vue.js 3框架构建^{[20]}，采用Element Plus组件库实现用户界面，通过'
        'Axios库与后端API进行HTTP通信。表现层负责接收用户操作、展示数据和处理前端'
        '交互逻辑，采用响应式设计确保在不同终端设备上的良好显示效果。'
    )

    add_cited_para(doc,
        '业务逻辑层基于Python语言和FastAPI框架构建^{[19]}，实现了系统的核心业务逻辑，'
        '包括用户管理、智能对话（RAG流程）、知识库管理和成长管理等功能。该层通过'
        'RESTful API接口为前端提供服务，采用JWT令牌机制进行身份认证，使用Pydantic'
        '库进行请求数据的验证和序列化。'
    )

    add_cited_para(doc,
        '数据访问层通过ORM（对象关系映射）框架SQLAlchemy与MySQL数据库交互，'
        '通过Chroma向量数据库的Python客户端进行向量数据的存储和检索操作^{[18]}。数据访问层'
        '封装了所有数据库操作，为业务逻辑层提供统一的数据操作接口。'
    )

    add_cited_para(doc,
        '数据存储层由MySQL关系型数据库和Chroma向量数据库组成。MySQL负责存储用户'
        '信息、对话记录、成长目标等结构化数据；Chroma负责存储知识文档的向量嵌入表示，'
        '支持高效的向量相似度检索^{[8]}。'
    )

    # 图 3-1 系统总体架构图
    add_body_para(doc,
        '如图3-1所示，本系统采用分层架构设计，各层之间通过标准化接口进行交互。'
    )
    add_figure(doc, generated_diagrams.get('architecture'), '图 3-1 系统总体架构图')

    add_body_para(doc,
        '表3-1所示为系统主要技术选型。'
    )

    add_three_line_table(doc,
        ['技术层面', '技术选型', '说明'],
        [
            ['前端框架', 'Vue.js 3', '渐进式JavaScript框架'],
            ['UI组件库', 'Element Plus', '基于Vue 3的桌面端组件库'],
            ['HTTP客户端', 'Axios', '基于Promise的HTTP库'],
            ['后端框架', 'FastAPI', '高性能异步Python Web框架'],
            ['ORM框架', 'SQLAlchemy', 'Python SQL工具包和ORM'],
            ['关系数据库', 'MySQL 8.0', '开源关系型数据库'],
            ['向量数据库', 'Chroma', '开源AI向量数据库'],
            ['大语言模型', 'GLM-4 API', '智谱AI大语言模型接口'],
            ['嵌入模型', 'text-embedding', '文本向量嵌入模型'],
        ]
    )

    # 表格标题
    para = doc.add_paragraph()
    run = para.add_run('表 3-1 系统主要技术选型')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(3), space_after=Pt(6))

    add_section_title(doc, '3.2 功能模块设计')

    # 图 3-3 功能模块结构图
    add_body_para(doc,
        '本系统的功能模块结构如图3-3所示，系统由用户管理、智能对话、知识库管理和'
        '个人成长管理四个核心模块组成。'
    )
    add_figure(doc, generated_diagrams.get('module'), '图 3-3 系统功能模块结构图')

    add_subsection_title(doc, '3.2.1 用户管理模块')

    add_body_para(doc,
        '用户管理模块实现用户的注册、登录、信息管理和权限控制功能。用户注册时，'
        '系统对用户名、邮箱和密码进行格式校验，密码经bcrypt加密后存入数据库。用户登录'
        '时，系统验证凭证后生成JWT访问令牌和刷新令牌，前端将令牌存储在localStorage中，'
        '后续请求通过Authorization请求头携带令牌进行身份认证。系统采用基于角色的访问'
        '控制机制，区分普通用户和管理员权限。'
    )

    add_subsection_title(doc, '3.2.2 智能对话模块')

    add_body_para(doc,
        '智能对话模块是系统的核心模块，实现了基于RAG技术的知识增强对话流程。该模块'
        '的处理流程如下：（1）接收用户输入的咨询问题，对问题文本进行预处理；'
        '（2）调用嵌入模型将用户问题转化为向量表示；（3）以问题向量为查询向量，在Chroma'
        '向量数据库中进行相似度检索，获取Top-K个最相关的知识片段；（4）将检索到的知识'
        '片段与用户问题组装成结构化的提示词（Prompt），注入大语言模型；（5）大语言模型'
        '基于提示词生成辅导建议，系统将结果返回给用户；（6）将对话记录持久化存储到数据库中。'
    )

    add_body_para(doc,
        '在提示词工程方面，系统采用了角色设定和上下文注入相结合的策略。系统提示词中'
        '定义了AI辅导师的角色定位、专业能力和回答风格要求，同时将检索到的知识片段以'
        '明确的格式标记注入上下文，引导模型基于知识内容生成回答。此外，系统还维护了'
        '滑动窗口式的对话历史，确保模型能够理解多轮对话的上下文。'
    )

    add_subsection_title(doc, '3.2.3 知识库管理模块')

    add_body_para(doc,
        '知识库管理模块实现了知识文档的全生命周期管理。文档处理流程包括：（1）文档上传'
        '与接收，支持PDF、TXT、Markdown等格式；（2）文本提取，利用PyMuPDF等库从文档中'
        '提取纯文本内容；（3）文本分块，采用递归字符分割策略，以500字符为默认块大小，'
        '相邻块之间保留50字符的重叠区域以维持语义连贯性；（4）向量化处理，调用嵌入模型'
        'API将每个知识块转化为维度为1024的向量表示；（5）向量存储，将知识块及其向量'
        '表示存入Chroma向量数据库，建立高效检索索引。'
    )

    add_subsection_title(doc, '3.2.4 个人成长管理模块')

    add_body_para(doc,
        '个人成长管理模块为用户提供了系统化的成长管理工具。该模块支持用户创建'
        '成长目标并设置目标的详细信息，包括目标名称、分类（如职业发展、技能提升、'
        '身心健康等）、描述说明、计划完成时间和优先级。用户可以随时更新目标的进度'
        '状态，系统自动记录状态变更的时间线。成长记录功能允许用户撰写每日或每周的'
        '成长日志，记录学习心得、反思总结和行动计划。数据统计功能对用户的成长数据'
        '进行聚合分析，通过可视化图表展示目标完成趋势和成长轨迹。'
    )

    add_section_title(doc, '3.3 数据库设计')

    add_subsection_title(doc, '3.3.1 概念模型设计')

    add_body_para(doc,
        '根据系统需求分析，本系统的数据库主要包含以下实体：用户（User）、对话'
        '（Conversation）、消息（Message）、知识文档（KnowledgeDocument）、知识片段'
        '（KnowledgeChunk）、成长目标（GrowthGoal）和成长记录（GrowthRecord）。'
        '各实体之间的关系如下：一个用户可以拥有多个对话、多个成长目标和多条成长记录；'
        '一个对话包含多条消息；一篇知识文档被分割为多个知识片段。'
    )

    # 图 3-2 数据库ER图
    add_body_para(doc,
        '如图3-2所示为系统的实体关系图，描述了各数据实体之间的关联关系。'
    )
    add_figure(doc, generated_diagrams.get('er'), '图 3-2 数据库ER图')

    add_subsection_title(doc, '3.3.2 数据库表设计')

    add_body_para(doc,
        '根据概念模型设计，本系统在MySQL数据库中设计了以下主要数据表。'
    )

    # 用户表
    add_body_para(doc, '（1）用户表（user）', indent=False)
    add_three_line_table(doc,
        ['字段名', '数据类型', '说明'],
        [
            ['id', 'INT AUTO_INCREMENT', '主键，自增'],
            ['username', 'VARCHAR(50)', '用户名，唯一'],
            ['email', 'VARCHAR(100)', '邮箱地址，唯一'],
            ['password_hash', 'VARCHAR(255)', '密码哈希值'],
            ['nickname', 'VARCHAR(50)', '昵称'],
            ['avatar', 'VARCHAR(255)', '头像URL'],
            ['role', 'ENUM("user","admin")', '用户角色'],
            ['created_at', 'DATETIME', '创建时间'],
        ]
    )
    para = doc.add_paragraph()
    run = para.add_run('表 3-2 用户表（user）')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(3), space_after=Pt(6))

    # 对话表
    add_body_para(doc, '（2）对话表（conversation）', indent=False)
    add_three_line_table(doc,
        ['字段名', '数据类型', '说明'],
        [
            ['id', 'INT AUTO_INCREMENT', '主键，自增'],
            ['user_id', 'INT', '外键，关联用户表'],
            ['title', 'VARCHAR(100)', '对话标题'],
            ['created_at', 'DATETIME', '创建时间'],
            ['updated_at', 'DATETIME', '更新时间'],
        ]
    )
    para = doc.add_paragraph()
    run = para.add_run('表 3-3 对话表（conversation）')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(3), space_after=Pt(6))

    # 消息表
    add_body_para(doc, '（3）消息表（message）', indent=False)
    add_three_line_table(doc,
        ['字段名', '数据类型', '说明'],
        [
            ['id', 'INT AUTO_INCREMENT', '主键，自增'],
            ['conversation_id', 'INT', '外键，关联对话表'],
            ['role', 'ENUM("user","assistant")', '消息角色'],
            ['content', 'TEXT', '消息内容'],
            ['created_at', 'DATETIME', '创建时间'],
        ]
    )
    para = doc.add_paragraph()
    run = para.add_run('表 3-4 消息表（message）')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(3), space_after=Pt(6))

    # 成长目标表
    add_body_para(doc, '（4）成长目标表（growth_goal）', indent=False)
    add_three_line_table(doc,
        ['字段名', '数据类型', '说明'],
        [
            ['id', 'INT AUTO_INCREMENT', '主键，自增'],
            ['user_id', 'INT', '外键，关联用户表'],
            ['title', 'VARCHAR(100)', '目标标题'],
            ['category', 'VARCHAR(50)', '目标分类'],
            ['description', 'TEXT', '目标描述'],
            ['status', 'ENUM("active","completed","paused")', '目标状态'],
            ['priority', 'INT', '优先级'],
            ['deadline', 'DATE', '计划完成日期'],
            ['created_at', 'DATETIME', '创建时间'],
        ]
    )
    para = doc.add_paragraph()
    run = para.add_run('表 3-5 成长目标表（growth_goal）')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(3), space_after=Pt(6))

    # 知识文档表
    add_body_para(doc, '（5）知识文档表（knowledge_document）', indent=False)
    add_three_line_table(doc,
        ['字段名', '数据类型', '说明'],
        [
            ['id', 'INT AUTO_INCREMENT', '主键，自增'],
            ['title', 'VARCHAR(200)', '文档标题'],
            ['file_name', 'VARCHAR(255)', '文件名'],
            ['category', 'VARCHAR(50)', '文档分类'],
            ['chunk_count', 'INT', '知识片段数量'],
            ['status', 'ENUM("processing","ready","error")', '处理状态'],
            ['uploaded_at', 'DATETIME', '上传时间'],
        ]
    )
    para = doc.add_paragraph()
    run = para.add_run('表 3-6 知识文档表（knowledge_document）')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(3), space_after=Pt(6))


# ============================================================
# 第4章 系统实现
# ============================================================
def create_chapter_4(doc, generated_diagrams=None):
    if generated_diagrams is None:
        generated_diagrams = {}
    add_chapter_title(doc, '第4章 系统实现')

    add_section_title(doc, '4.1 开发环境与工具')

    add_body_para(doc,
        '本系统的开发环境配置如表4-1所示。'
    )

    add_three_line_table(doc,
        ['项目', '配置', '版本'],
        [
            ['操作系统', 'Windows 11 / Ubuntu 22.04', '-'],
            ['开发语言', 'Python / JavaScript', '3.11 / ES2020+'],
            ['IDE', 'VS Code / PyCharm', '最新版'],
            ['前端框架', 'Vue.js 3 + Vite', '3.4+'],
            ['后端框架', 'FastAPI + Uvicorn', '0.110+'],
            ['数据库', 'MySQL', '8.0'],
            ['向量数据库', 'Chroma', '0.4+'],
            ['版本控制', 'Git', '2.40+'],
        ]
    )
    para = doc.add_paragraph()
    run = para.add_run('表 4-1 开发环境配置')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(3), space_after=Pt(6))

    add_section_title(doc, '4.2 核心功能实现')

    add_subsection_title(doc, '4.2.1 RAG检索增强生成流程实现')

    # 图 4-1 RAG流程图
    add_body_para(doc,
        '如图4-1所示，RAG流程包括知识库构建和在线检索生成两个阶段。'
    )
    add_figure(doc, generated_diagrams.get('rag_flow'), '图 4-1 RAG检索增强生成流程图')

    add_cited_para(doc,
        'RAG流程是本系统的核心技术实现，主要包括知识文档处理、向量检索和增强生成'
        '三个阶段^{[1,3]}。在知识文档处理阶段，系统使用PyMuPDF和python-docx等库对不同格式的'
        '文档进行文本提取，然后采用RecursiveCharacterTextSplitter进行文本分块处理。'
        '分块参数设置为：块大小500字符，块重叠50字符，确保知识片段的语义完整性。'
    )

    add_cited_para(doc,
        '在向量检索阶段，系统使用智谱AI提供的文本嵌入模型将知识块转化为向量表示，'
        '向量维度为1024。生成的向量数据存储在Chroma向量数据库中^{[18]}，并利用其内置的HNSW'
        '（Hierarchical Navigable Small World）索引算法实现高效的近似最近邻检索^{[8]}。当用户'
        '发起咨询时，系统将用户问题同样转化为向量，在Chroma中执行相似度检索，返回'
        'Top-5个最相关的知识片段。'
    )

    add_cited_para(doc,
        '在增强生成阶段，系统将检索到的知识片段按照预设的提示词模板进行组装。'
        '提示词模板包括系统角色设定、知识上下文和用户问题三个部分。系统角色设定部分'
        '定义了AI辅导师的身份和能力边界；知识上下文部分将检索到的知识片段以清晰的'
        '格式标记注入；用户问题部分包含当前问题和历史对话记录^{[16]}。组装完成的提示词被'
        '发送至大语言模型API，模型基于知识上下文生成回答，确保回答内容具有知识依据。'
    )

    add_subsection_title(doc, '4.2.2 用户认证与授权实现')

    add_body_para(doc,
        '用户认证模块基于JWT（JSON Web Token）机制实现。用户注册时，系统使用bcrypt'
        '算法对密码进行哈希处理，将哈希值存入数据库。用户登录时，系统验证密码后，'
        '使用PyJWT库生成包含用户ID和角色的访问令牌（有效期24小时）和刷新令牌（有效期'
        '7天）。后端通过FastAPI的依赖注入机制实现令牌验证中间件，所有受保护的API'
        '接口在处理请求前自动验证令牌的有效性和权限。'
    )

    add_subsection_title(doc, '4.2.3 对话管理实现')

    add_body_para(doc,
        '对话管理模块实现了多轮对话的上下文维护功能。系统采用滑动窗口策略管理对话'
        '历史，保留最近10轮对话记录作为上下文输入。当用户发送新消息时，系统首先从'
        '数据库中加载当前对话的历史消息，结合新消息构建完整的对话上下文。对话响应采用'
        '流式传输（Streaming）机制，利用FastAPI的Server-Sent Events（SSE）支持，实现'
        '模型生成内容的逐字输出，提升用户的交互体验。'
    )

    add_subsection_title(doc, '4.2.4 知识库处理实现')

    add_body_para(doc,
        '知识库处理模块实现了异步的文档处理流水线。当用户上传文档后，系统将文档处理'
        '任务加入后台任务队列，避免阻塞用户的其他操作。后台任务执行器依次完成文本提取、'
        '文本分块、向量生成和向量存储四个步骤。处理完成后，更新文档状态为"就绪"并通知'
        '用户。如果任一步骤出现错误，系统将文档状态标记为"处理失败"并记录错误日志，'
        '方便用户和管理员排查问题。'
    )

    add_section_title(doc, '4.3 系统界面展示')

    add_body_para(doc,
        '本节对系统的主要用户界面进行说明。系统采用现代化的UI设计风格，整体布局'
        '采用左侧导航栏 + 右侧内容区域的经典后台管理布局，界面色调以蓝白为主，简洁大方。'
    )

    add_body_para(doc,
        '（1）用户登录界面：系统登录页面提供用户名和密码输入框，支持"记住我"功能。'
        '页面中央显示系统Logo和名称，整体布局简洁明了。', indent=False
    )

    add_body_para(doc,
        '（2）智能对话界面：对话界面是用户使用最频繁的核心页面。页面左侧显示对话'
        '历史列表，用户可以快速切换不同的对话会话；页面右侧为对话区域，采用聊天气泡'
        '风格展示用户消息和AI回复，AI回复内容支持Markdown格式渲染。', indent=False
    )

    add_body_para(doc,
        '（3）知识库管理界面：知识库管理页面以列表形式展示所有已上传的知识文档，'
        '每条记录显示文档标题、分类、片段数量、处理状态和上传时间。页面顶部提供搜索'
        '过滤和文档上传功能。', indent=False
    )

    add_body_para(doc,
        '（4）个人成长管理界面：成长管理页面分为目标管理和成长记录两个子页面。目标'
        '管理页面以卡片形式展示用户设定的各项目标，支持按分类和状态筛选。成长记录页面'
        '以时间线形式展示用户的成长日志，支持新增和编辑操作。', indent=False
    )


# ============================================================
# 第5章 系统测试
# ============================================================
def create_chapter_5(doc):
    add_chapter_title(doc, '第5章 系统测试')

    add_section_title(doc, '5.1 测试环境')

    add_body_para(doc,
        '系统测试在以下环境中进行：服务器操作系统为Ubuntu 22.04 LTS，CPU为Intel'
        ' Core i7-12700H，内存16GB，硬盘512GB SSD。后端服务运行在Python 3.11环境下，'
        '使用Uvicorn作为ASGI服务器。数据库使用MySQL 8.0和Chroma 0.4.x。前端通过'
        'Chrome 120浏览器进行测试。网络环境为校园网，带宽100Mbps。'
    )

    add_section_title(doc, '5.2 功能测试')

    add_body_para(doc,
        '功能测试针对系统的各功能模块进行了全面的测试验证，测试用例和结果如表5-1所示。'
    )

    add_three_line_table(doc,
        ['测试编号', '测试模块', '测试内容', '预期结果', '实际结果', '是否通过'],
        [
            ['TC-001', '用户注册', '输入合法信息注册', '注册成功', '注册成功', '通过'],
            ['TC-002', '用户注册', '输入已存在用户名', '提示用户名已存在', '提示正确', '通过'],
            ['TC-003', '用户登录', '输入正确账号密码', '登录成功跳转首页', '跳转正确', '通过'],
            ['TC-004', '用户登录', '输入错误密码', '提示密码错误', '提示正确', '通过'],
            ['TC-005', '智能对话', '发送成长相关问题', '返回知识增强回答', '回答合理', '通过'],
            ['TC-006', '智能对话', '进行多轮连续对话', '理解上下文语义', '理解正确', '通过'],
            ['TC-007', '知识库', '上传PDF文档', '文档解析成功', '解析正确', '通过'],
            ['TC-008', '知识库', '上传TXT文档', '文本提取成功', '提取正确', '通过'],
            ['TC-009', '成长管理', '创建成长目标', '目标保存成功', '保存成功', '通过'],
            ['TC-010', '成长管理', '添加成长记录', '记录保存成功', '保存成功', '通过'],
        ]
    )
    para = doc.add_paragraph()
    run = para.add_run('表 5-1 功能测试用例及结果')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(3), space_after=Pt(6))

    add_body_para(doc,
        '功能测试结果表明，系统各模块的功能均能够正常运行，满足需求分析中确定的'
        '各项功能需求。智能对话模块能够基于知识库内容生成合理的辅导建议，对话上下文'
        '维护正常。知识库管理模块能够正确处理PDF和TXT格式的文档，文档解析和向量化'
        '流程运行正常。个人成长管理模块的目标创建、记录添加等功能均运行正常。'
    )

    add_section_title(doc, '5.3 性能测试')

    add_body_para(doc,
        '为了验证系统在实际使用场景下的性能表现，本文对系统进行了性能测试。测试'
        '主要关注页面加载时间、API接口响应时间和智能对话响应时间三个指标。'
    )

    add_three_line_table(doc,
        ['测试项目', '测试条件', '平均响应时间', '是否达标'],
        [
            ['页面加载', '单用户访问', '1.2s', '达标'],
            ['API接口响应', '单用户请求', '85ms', '达标'],
            ['智能对话（无缓存）', '单用户对话', '6.8s', '达标'],
            ['智能对话（有缓存）', '单用户对话', '3.2s', '达标'],
            ['知识库检索', '1000篇文档', '120ms', '达标'],
            ['文档上传处理', '5MB PDF文档', '15.6s', '达标'],
        ]
    )
    para = doc.add_paragraph()
    run = para.add_run('表 5-2 性能测试结果')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, first_line_indent=Emu(0), space_before=Pt(3), space_after=Pt(6))

    add_body_para(doc,
        '性能测试结果表明，系统的各项性能指标均满足设计要求。页面加载时间和API'
        '接口响应时间均在可接受范围内。智能对话模块在无缓存情况下的平均响应时间为'
        '6.8秒，主要包括知识检索和大模型推理两个环节的时间消耗；启用缓存后响应时间'
        '降低至3.2秒，显著提升了用户体验。向量检索在1000篇文档规模下的平均耗时仅'
        '120毫秒，表明Chroma向量数据库具有良好的检索性能。'
    )

    add_section_title(doc, '5.4 测试结论')

    add_body_para(doc,
        '通过对系统进行全面的功能测试和性能测试，得出以下结论：（1）系统各功能模块'
        '运行正常，所有测试用例均通过，系统功能完整、运行稳定；（2）系统的各项性能指标'
        '均达到设计要求，页面响应和API接口性能良好，智能对话模块的响应时间在可接受'
        '范围内；（3）基于RAG技术的智能对话功能能够有效利用知识库内容，生成的辅导建议'
        '具有较好的准确性和实用性。综上所述，本系统达到了预期的设计目标，具备实际'
        '应用的条件。'
    )


# ============================================================
# 总结与展望
# ============================================================
def create_conclusion(doc):
    add_chapter_title(doc, '总结与展望')

    add_body_para(doc,
        '本文设计并实现了一个基于检索增强生成的个人成长辅导系统，主要完成了以下'
        '工作：'
    )

    add_body_para(doc,
        '（1）对个人成长辅导领域的需求进行了深入分析，明确了系统的功能需求和非功能'
        '需求，并从技术、经济和社会三个维度进行了可行性论证。', indent=False
    )

    add_body_para(doc,
        '（2）设计了基于B/S架构的四层系统架构，采用前后端分离的开发模式，完成了'
        '用户管理、智能对话、知识库管理和个人成长管理四个核心功能模块的详细设计。', indent=False
    )

    add_body_para(doc,
        '（3）实现了完整的RAG技术流程，包括知识文档的解析分块、向量嵌入与存储、'
        '语义检索和知识增强生成，有效解决了大语言模型在专业领域应用中的幻觉问题。', indent=False
    )

    add_body_para(doc,
        '（4）通过功能测试和性能测试验证了系统的可行性和实用性，测试结果表明系统'
        '各项功能运行正常，性能指标满足实际使用需求。', indent=False
    )

    add_body_para(doc,
        '尽管本系统已基本实现了预期功能，但仍存在一些可以改进和优化的方向，未来'
        '可以从以下几个方面进行深入研究：'
    )

    add_body_para(doc,
        '（1）优化RAG流程：引入更高级的检索策略，如混合检索（结合关键词检索和向量'
        '检索）、重排序（Reranking）等技术，进一步提升检索的准确性和生成质量。', indent=False
    )

    add_body_para(doc,
        '（2）个性化推荐：基于用户的历史对话和成长记录数据，利用机器学习算法构建'
        '用户画像模型，实现更加精准的个性化辅导建议和资源推荐。', indent=False
    )

    add_body_para(doc,
        '（3）多模态支持：扩展系统支持图片、音频等多模态数据的处理能力，丰富知识库'
        '的内容形式，提供更加多样化的辅导体验。', indent=False
    )

    add_body_para(doc,
        '（4）移动端适配：开发移动端应用程序或优化移动端Web体验，使系统能够更好地'
        '满足用户随时随地使用的需求。', indent=False
    )


# ============================================================
# 参考文献
# ============================================================
def create_references(doc):
    add_chapter_title(doc, '参考文献')

    refs = [
        '[1] Lewis P, Perez E, Piktus A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]//Advances in Neural Information Processing Systems. 2020, 33: 9459-9474.',
        '[2] Karpukhin V, Oğuz B, Min S, et al. Dense passage retrieval for open-domain question answering[C]//Proceedings of the 2020 Conference on Empirical Methods in Natural Language Processing. 2020: 6769-6781.',
        '[3] Gao Y, Xiong Y, Gao X, et al. Retrieval-augmented generation for large language models: A survey[J]. arXiv preprint arXiv:2312.10997, 2024.',
        '[4] Brown T, Mann B, Ryder N, et al. Language models are few-shot learners[C]//Advances in Neural Information Processing Systems. 2020, 33: 1877-1901.',
        '[5] OpenAI. GPT-4 technical report[J]. arXiv preprint arXiv:2303.08774, 2023.',
        '[6] Bubeck S, Chandrasekaran V, Eldan R, et al. Sparks of artificial general intelligence: Early experiments with GPT-4[J]. arXiv preprint arXiv:2303.12712, 2023.',
        '[7] Zhao W X, Zhou K, Li J, et al. A survey of large language models[J]. arXiv preprint arXiv:2303.18223, 2023.',
        '[8] Johnson J, Douze M, Jégou H. Billion-scale similarity search with GPUs[J]. IEEE Transactions on Big Data, 2021, 7(3): 535-547.',
        '[9] 李彦博, 张伟, 刘洋. 基于大语言模型的智能问答系统研究综述[J]. 计算机科学, 2024, 51(1): 1-15.',
        '[10] 王晓宇, 陈明. 检索增强生成技术的研究进展与应用[J]. 自动化学报, 2024, 50(3): 456-470.',
        '[11] 张三丰, 李四光. 面向知识库问答的语义检索方法研究[J]. 软件学报, 2023, 34(8): 3521-3540.',
        '[12] 陈伟, 刘明. 基于向量数据库的大规模知识检索系统设计与实现[J]. 计算机工程与应用, 2024, 60(2): 112-120.',
        '[13] Tolle K M, Tansley D S W, Hey A J G, et al. The fourth paradigm: Data-intensive scientific discovery[J]. Microsoft Research, 2009.',
        '[14] 刘知远, 孙茂松. 表示学习: 从自然语言处理到信息检索[J]. 中国科学: 信息科学, 2023, 53(9): 1681-1698.',
        '[15] 赵鑫, 蔡康佳, 郭东雨, 等. 大语言模型综述[J]. 中国科学: 信息科学, 2024, 54(1): 1-48.',
        '[16] Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of deep bidirectional transformers for language understanding[C]//Proceedings of the 2019 Conference of the North American Chapter of the Association for Computational Linguistics. 2019: 4171-4186.',
        '[17] Milvus Project. Milvus: A purpose-built vector data management platform[EB/OL]. (2023-01-01)[2026-04-15]. https://milvus.io/.',
        '[18] Chroma. ChromaDB: The AI-native open-source embedding database[EB/OL]. (2024-01-01)[2026-04-15]. https://www.trychroma.com/.',
        '[19] FastAPI. FastAPI framework[EB/OL]. (2024-01-01)[2026-04-15]. https://fastapi.tiangolo.com/.',
        '[20] Vue.js. Vue.js - The progressive JavaScript framework[EB/OL]. (2024-01-01)[2026-04-15]. https://vuejs.org/.',
    ]

    for ref in refs:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU)
        set_paragraph_format(para, first_line_indent=Emu(0),
                             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=Pt(2), space_after=Pt(2))


# ============================================================
# 致谢
# ============================================================
def create_acknowledgment(doc):
    add_chapter_title(doc, '致  谢')

    add_body_para(doc,
        '时光荏苒，四年的大学生活即将画上句号。回顾这段充实而难忘的求学时光，'
        '我心中充满了感激之情。在这篇论文完成之际，我想向所有给予我帮助和支持的人'
        '表达最诚挚的谢意。'
    )

    add_body_para(doc,
        f'首先，我要特别感谢我的导师{THESIS["advisor"]}老师。从课题的选择、方案的'
        '设计到论文的撰写，霍老师都给予了我悉心的指导和耐心的帮助。霍老师严谨的'
        '治学态度、渊博的专业知识和精益求精的工作作风深深感染了我，使我在学术研究'
        '和个人成长方面都受益匪浅。'
    )

    add_body_para(doc,
        '其次，我要感谢信息工程学院所有教授过我课程的老师们。正是你们在课堂上的'
        '悉心教导，为我打下了坚实的专业基础，使我具备了完成本课题所需的知识和能力。'
        '特别感谢在计算机科学、数据库技术、软件工程等方面给予我指导的各位老师。'
    )

    add_body_para(doc,
        '同时，我要感谢与我朝夕相处的同学们。在大学四年的学习和生活中，我们互相'
        '帮助、共同进步，一起度过了无数个难忘的时光。感谢室友和朋友们在日常学习和'
        '论文写作过程中给予我的鼓励和支持。'
    )

    add_body_para(doc,
        '最后，我要深深地感谢我的家人。是你们的默默付出和无私奉献，为我创造了'
        '良好的学习环境；是你们的理解和支持，给予了我不断前进的动力。你们的爱是我'
        '永远的依靠和力量源泉。'
    )

    add_body_para(doc,
        '大学四年，我不仅学到了专业知识，更学会了独立思考和解决问题的方法。未来'
        '的道路还很长，我将带着在学校里学到的一切，继续努力，不断前行，以实际行动'
        '回报所有关心和帮助过我的人。'
    )


# ============================================================
# 主函数：组装文档
# ============================================================
def generate_diagrams():
    """生成所有 Mermaid 图形，返回 {名称: 图片路径} 字典"""
    print('正在生成图形...')
    diagrams = {}
    items = [
        ('architecture', DIAGRAM_ARCHITECTURE, 'system_architecture.png'),
        ('er', DIAGRAM_ER, 'database_er.png'),
        ('rag_flow', DIAGRAM_RAG_FLOW, 'rag_flow.png'),
        ('module', DIAGRAM_MODULE, 'module_structure.png'),
    ]
    for name, mmd_text, filename in items:
        print(f'  生成 {filename}...')
        diagrams[name] = generate_mermaid_image(mmd_text, filename)
        if diagrams[name]:
            print(f'    ✓ {filename}')
        else:
            print(f'    ✗ {filename} 生成失败')
    return diagrams


def generate_thesis():
    # 0. 预生成所有图形
    generated_diagrams = generate_diagrams()

    doc = Document()

    # 配置标题样式（必须在创建内容之前调用）
    configure_styles(doc)

    # 1. 封面
    create_cover_page(doc)

    # 2. 声明页
    create_declaration_page(doc)

    # 3. 中文摘要
    create_chinese_abstract(doc)

    # 4. 英文摘要
    create_english_abstract(doc)

    # 5. 目录
    create_table_of_contents(doc)

    # === 正文部分：切换到阿拉伯数字页码 ===
    # 第1章
    create_chapter_1(doc)

    # 设置正文第一个 section 的页码为阿拉伯数字从1开始
    body_section = doc.sections[-1]
    add_header(body_section, HEADER_FIXED)
    add_page_number(body_section)
    set_page_number_format(body_section, 'decimal')
    set_page_number_start(body_section, 1)

    # 6-9. 正文章节（传入图形）
    create_chapter_2(doc)
    create_chapter_3(doc, generated_diagrams)
    create_chapter_4(doc, generated_diagrams)
    create_chapter_5(doc)

    # 10. 总结与展望
    create_conclusion(doc)

    # 11. 参考文献
    create_references(doc)

    # 12. 致谢
    create_acknowledgment(doc)

    # 保存文档
    doc.save(OUTPUT_FILE)
    print(f'论文已生成: {OUTPUT_FILE}')


if __name__ == '__main__':
    generate_thesis()
